import subprocess
import sys
import os

def build_docx_report():
    # Ensure python-docx is installed
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn, nsdecls

    doc = Document()

    # Define Colors
    COLOR_PRIMARY = RGBColor(13, 17, 23)     # Dark slate
    COLOR_SECONDARY = RGBColor(59, 130, 246) # Accent Blue
    COLOR_MUTED = RGBColor(139, 148, 158)    # Muted Grey

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ---------------------------------------------------------
    # Helper functions for styling
    # ---------------------------------------------------------
    def set_font(run, name="Calibri", size_pt=11, color=None, bold=False, italic=False):
        run.font.name = name
        run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = color
        run.bold = bold
        run.italic = italic

    def set_cell_background(cell, color_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_table_borders(table):
        tblPr = table._tbl.tblPr
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>\n'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>\n'
            '  <w:left w:val="none"/>\n'
            '  <w:right w:val="none"/>\n'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>\n'
            '  <w:insideV w:val="none"/>\n'
            '</w:tblBorders>'
        )
        tblPr.append(tblBorders)

    def add_page_number(run):
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    # ---------------------------------------------------------
    # PAGE 1: TITLE PAGE
    # ---------------------------------------------------------
    # Spacing for alignment
    for _ in range(3):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("UNIVERSAL CLOUD PRINT SYSTEM\n(UCPS)")
    set_font(r_title, "Arial", 26, COLOR_PRIMARY, bold=True)
    p_title.paragraph_format.space_after = Pt(12)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("A Secure, QR-Based Smart Cloud Printing Platform for Legacy Printers")
    set_font(r_sub, "Arial", 14, COLOR_SECONDARY, italic=True)
    p_sub.paragraph_format.space_after = Pt(40)

    p_sub_details = doc.add_paragraph()
    p_sub_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub_details = p_sub_details.add_run(
        "Submitted To\n"
        "CSE Fest 2026 Software Project Showcase\n"
        "Department of Computer Science and Engineering (CSE)\n"
        "Shanto-Mariam University of Creative Technology (SMUCT)\n"
    )
    set_font(r_sub_details, "Calibri", 12, bold=False)
    p_sub_details.paragraph_format.space_after = Pt(30)

    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_by = p_by.add_run("Submitted By")
    set_font(r_by, "Calibri", 12, COLOR_MUTED, bold=True)
    p_by.paragraph_format.space_after = Pt(10)

    # Team Members Table
    table_team = doc.add_table(rows=4, cols=3)
    table_team.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_team)

    headers = ["Name", "Student ID", "Role"]
    col_widths = [Inches(2.5), Inches(1.8), Inches(1.8)]

    # Header Row
    for idx, text in enumerate(headers):
        cell = table_team.rows[0].cells[idx]
        cell.width = col_widths[idx]
        set_cell_background(cell, "0D1117")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, "Arial", 10, RGBColor(255, 255, 255), bold=True)

    team_data = [
        ["Wahidur Rahman", "2026-2-60-046", "Team Leader / Full-Stack"],
        ["Maloy Roy Orko", "2026-2-60-462", "Backend & Python Developer"],
        ["Intisar Muhib", "2026-2-60-239", "Frontend Designer / QA"]
    ]

    for row_idx, row_data in enumerate(team_data):
        row = table_team.rows[row_idx + 1]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, "Calibri", 10)

    doc.add_paragraph().paragraph_format.space_before = Pt(30)
    
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_inst = p_inst.add_run(
        "Department of Computer Science & Engineering (CSE)\n"
        "East West University (EWU)\n"
        "Date of Submission: June 2026"
    )
    set_font(r_inst, "Calibri", 11, bold=True)

    # ---------------------------------------------------------
    # PAGE 2: TABLE OF CONTENTS
    # ---------------------------------------------------------
    doc.add_page_break()
    p_h = doc.add_paragraph()
    r_h = p_h.add_run("Table of Contents")
    set_font(r_h, "Arial", 18, COLOR_PRIMARY, bold=True)
    p_h.paragraph_format.space_after = Pt(20)

    toc_items = [
        ("Section 1: Title Page", "1"),
        ("Section 2: Abstract", "3"),
        ("Section 3: Problem Statement", "4"),
        ("Section 4: Objectives", "4"),
        ("Section 5: System Architecture", "5"),
        ("Section 6: Technologies Used", "7"),
        ("Section 7: Key Features", "8"),
        ("Section 8: Implementation Details", "9"),
        ("Section 9: Results & Evaluation", "11"),
        ("Section 10: Challenges & Limitations", "13"),
        ("Section 11: Future Scope", "13"),
        ("Section 12: References", "14")
    ]

    for item, page in toc_items:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_after = Pt(6)
        r_item = p_item.add_run(item)
        set_font(r_item, "Calibri", 11)
        r_dots = p_item.add_run(" " + "." * (80 - len(item)) + " ")
        set_font(r_dots, "Calibri", 11, COLOR_MUTED)
        r_page = p_item.add_run(page)
        set_font(r_page, "Calibri", 11, bold=True)

    # ---------------------------------------------------------
    # SECTION 2: ABSTRACT
    # ---------------------------------------------------------
    doc.add_page_break()
    p_sec = doc.add_paragraph()
    r_sec = p_sec.add_run("Section 2: Abstract")
    set_font(r_sec, "Arial", 16, COLOR_PRIMARY, bold=True)
    p_sec.paragraph_format.space_after = Pt(12)

    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.line_spacing = 1.15
    p_abs.paragraph_format.space_after = Pt(12)
    r_abs = p_abs.add_run(
        "In modern academic and commercial environments, document printing remains a core administrative and "
        "educational necessity. Despite rapid advances in cloud computing, file transfer, and web technologies, "
        "the workflow surrounding physical document printing remains highly archaic, manual, and fragmented. "
        "Users routinely transfer files to print operators using temporary channels such as WhatsApp, Gmail, "
        "Facebook Messenger, Bluetooth, or physical USB storage devices. These methods introduce substantial "
        "operational bottlenecks, extend waiting times, compromise document privacy, and require high manual effort "
        "from printing staff.\n\n"
        "The Universal Cloud Print System (UCPS) is an innovative, centralized, cloud-based platform designed "
        "to modernize legacy printing services. UCPS enables users to securely upload documents from any "
        "internet-connected device (smartphone, tablet, laptop, or desktop) and print them on any registered physical "
        "printer via a quick QR code pairing mechanism. Operating as an intelligent middleware with a zero-configuration "
        "SQLite database engine, UCPS coordinates data flow between the cloud interface and a custom Tkinter-based "
        "GUI Desktop Print Client running on local print nodes. The desktop application polls the secure server, "
        "downloads buffer files, and spools them automatically to local spoolers, eliminating third-party messaging "
        "dependencies and security vulnerabilities.\n\n"
        "To protect user privacy and system integrity, the platform implements a secure, multi-stage file validation "
        "pipeline. Uploaded files undergo binary MIME type verification, strict size constraint validation (limited to 10MB), "
        "and SHA-256 filename hashing. The backend isolates files inside a protected directory secured by custom "
        ".htaccess rules on Apache/cPanel systems. Jobs are processed using a First-In-First-Out (FIFO) queue scheduler, "
        "and files are automatically deleted from the server immediately after successful printing, ensuring compliance "
        "with modern data privacy standards.\n\n"
        "Quantitative evaluation in a simulated university campus environment demonstrated an average 60% reduction in "
        "user steps (from 5 to 2), an 83% reduction in operator actions (from 6 to 1), and a substantial decrease in average "
        "document transfer time (from 120–180 seconds down to 15–30 seconds). UCPS is a highly viable, secure, and "
        "cost-effective solution suitable for academic institutions, libraries, and commercial print shops seeking to "
        "transition to smart, digital-first printing workflows."
    )
    set_font(r_abs, "Calibri", 11)

    p_kw = doc.add_paragraph()
    r_kw_title = p_kw.add_run("Keywords: ")
    set_font(r_kw_title, "Calibri", 11, bold=True)
    r_kw_text = p_kw.add_run("Cloud Printing, Smart Spooler, GUI Desktop Client, SQLite Database, QR Code Pairing, Document Security.")
    set_font(r_kw_text, "Calibri", 11, italic=True)

    # ---------------------------------------------------------
    # SECTION 3: PROBLEM STATEMENT
    # ---------------------------------------------------------
    doc.add_page_break()
    p_sec = doc.add_paragraph()
    r_sec = p_sec.add_run("Section 3: Problem Statement")
    set_font(r_sec, "Arial", 16, COLOR_PRIMARY, bold=True)
    p_sec.paragraph_format.space_after = Pt(12)

    p_prob = doc.add_paragraph()
    p_prob.paragraph_format.line_spacing = 1.15
    p_prob.paragraph_format.space_after = Pt(12)
    r_prob = p_prob.add_run(
        "Modern document printing workflows remain inefficient and insecure, relying on manual processes. "
        "In university campus centers and commercial copy shops, users must send documents via temporary communication "
        "channels such as WhatsApp, Gmail, Messenger, or USB drives to print operators. This workflow introduces "
        "several problems:\n\n"
        "1. Privacy Violations: Personal files (NID cards, transcripts, certificates, bank statements) remain "
        "permanently in print operators' messaging histories, chat windows, and download directories on public computers, "
        "exposing sensitive data to unauthorized recovery.\n"
        "2. Operational Bottlenecks: Operators must manually download each file, open the corresponding application "
        "(Acrobat Reader, Word, Photos), adjust print configurations, select local spoolers, and execute prints. "
        "This manual process takes 6 discrete steps, causing significant delays during peak hours.\n"
        "3. Lack of Queue Visibility: Users cannot monitor print queues or printer statuses remotely. They must wait "
        "physically in front of the print counter to find out if a printer is offline, busy, or running out of paper/ink.\n"
        "4. High Implementation Cost: Existing smart cloud printing hardware (like ePrint-enabled printers) is "
        "expensive and requires replacing existing legacy printers, which is financially impractical for academic "
        "environments and small business owners."
    )
    set_font(r_prob, "Calibri", 11)

    # ---------------------------------------------------------
    # SECTION 4: OBJECTIVES
    # ---------------------------------------------------------
    p_sec = doc.add_paragraph()
    p_sec.paragraph_format.space_before = Pt(24)
    r_sec = p_sec.add_run("Section 4: Objectives")
    set_font(r_sec, "Arial", 16, COLOR_PRIMARY, bold=True)
    p_sec.paragraph_format.space_after = Pt(12)

    p_obj_intro = doc.add_paragraph()
    r_obj_intro = p_obj_intro.add_run(
        "The primary goal of UCPS is to develop a software-only cloud printing platform that modernizes legacy physical "
        "printers. The specific engineering objectives are:"
    )
    set_font(r_obj_intro, "Calibri", 11)

    objs = [
        ("Direct Web Portal Upload:", " Provide a secure web portal where users can register, upload documents directly, select printer nodes, and track spooling status in real-time."),
        ("QR-Based Pairing Interface:", " Deploy dynamic QR code labels on physical printer bodies that instantly bind a user's web browser session with the selected printer node on scan."),
        ("Desktop Client with Operator GUI:", " Build a secure, dark-themed Tkinter GUI desktop client executable (UCPS_PrintNode.exe) featuring operator authentication, active spooler logs mirroring, and dynamic printer auto-detection."),
        ("Integrated Cash and Online Billing:", " Integrate tokenized online checkout (e.g. bKash checkout API) along with a manual Cash Payment approval model in the Operator dashboard."),
        ("Asynchronous FIFO Print Scheduler:", " Develop a database-backed First-In-First-Out scheduling model executing atomic transaction queue locks to eliminate race conditions."),
        ("Zero-Configuration Database:", " Implement SQLite 3 with PDO connection fallbacks, ensuring compatibility across local XAMPP setups and standard shared cPanel hosting servers."),
        ("Automated Document Disposal:", " Ensure complete privacy by automatically deleting temporary files from server disk buffers immediately after spooler confirmation.")
    ]

    for title, desc in objs:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.paragraph_format.space_after = Pt(4)
        r_bt = p_bullet.add_run(title)
        set_font(r_bt, "Calibri", 11, bold=True)
        r_bd = p_bullet.add_run(desc)
        set_font(r_bd, "Calibri", 11)

    # ---------------------------------------------------------
    # SECTION 5: SYSTEM ARCHITECTURE
    # ---------------------------------------------------------
    doc.add_page_break()
    p_sec = doc.add_paragraph()
    r_sec = p_sec.add_run("Section 5: System Architecture")
    set_font(r_sec, "Arial", 16, COLOR_PRIMARY, bold=True)
    p_sec.paragraph_format.space_after = Pt(12)

    p_arch = doc.add_paragraph()
    r_arch = p_arch.add_run(
        "The system runs on a highly decoupled Client-Server-Daemon architecture. The Client Web Portal runs "
        "inside standard mobile/desktop browsers. The Web Server manages requests, file storage, and queue operations, "
        "communicating with a secure SQLite database. The local Execution Client (Print Node) operates directly on a "
        "local computer physically connected to printers, executing print tasks and feeding live logs back to the server."
    )
    set_font(r_arch, "Calibri", 11)
    p_arch.paragraph_format.space_after = Pt(12)

    # System Architecture Table
    table_arch = doc.add_table(rows=4, cols=3)
    table_arch.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_arch)

    arch_headers = ["System Layer", "Key Components / Modules", "Role & Protocols"]
    arch_col_widths = [Inches(1.8), Inches(2.4), Inches(2.3)]

    for idx, text in enumerate(arch_headers):
        cell = table_arch.rows[0].cells[idx]
        cell.width = arch_col_widths[idx]
        set_cell_background(cell, "0D1117")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, "Arial", 10, RGBColor(255, 255, 255), bold=True)

    arch_data = [
        ["Presentation Layer (Client Browser)", "Web Browser UI (HTML5/CSS3/JS), Web QR Scanner, User Dashboard, Operator Spooler Panel.", "Provides the interface for login, pairing, cost checking, payment, and file submission. Communicates via HTTPS JSON APIs."],
        ["Application Layer (Web Server & DB)", "PHP 8.x Backend APIs, SQLite 3 Database (ucps_db.sqlite), Upload Validator, CORS Middleware, .htaccess Rules.", "Validates file MIME-types, hashes filenames, calculates costs, tracks transactions, queues jobs, and restricts direct database access."],
        ["Execution Layer (Print Node Daemon)", "Tkinter GUI Operator Panel (UCPS_PrintNode.exe), Background Queue Thread, OS Spooler (SumatraPDF on Windows, CUPS on Linux).", "Logs in operator, detects system printers, pushes printer states to DB, polls for pending print jobs, spools files silently, and triggers post-print status updates."]
    ]

    for row_idx, row_data in enumerate(arch_data):
        row = table_arch.rows[row_idx + 1]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = arch_col_widths[col_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, "Calibri", 10)

    # ---------------------------------------------------------
    # DATA FLOW DESIGN
    # ---------------------------------------------------------
    p_dfd_h = doc.add_paragraph()
    p_dfd_h.paragraph_format.space_before = Pt(24)
    r_dfd_h = p_dfd_h.add_run("5.2 Data Flow Diagram (DFD) Mapping")
    set_font(r_dfd_h, "Arial", 13, COLOR_PRIMARY, bold=True)
    p_dfd_h.paragraph_format.space_after = Pt(10)

    # Level-0 Context Table
    p_l0 = doc.add_paragraph()
    r_l0 = p_l0.add_run("Level-0 DFD Table (System Context Diagram)")
    set_font(r_l0, "Calibri", 11, bold=True)
    
    table_l0 = doc.add_table(rows=3, cols=3)
    table_l0.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_l0)
    
    l0_headers = ["Source Entity", "Data Input to UCPS", "Data Output from UCPS"]
    l0_widths = [Inches(1.8), Inches(2.3), Inches(2.4)]
    
    for idx, text in enumerate(l0_headers):
        cell = table_l0.rows[0].cells[idx]
        cell.width = l0_widths[idx]
        set_cell_background(cell, "0D1117")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, "Arial", 10, RGBColor(255, 255, 255), bold=True)

    l0_data = [
        ["User / Student Client", "Login details, QR Printer ID, Printable documents (.pdf, .png, .jpg), Payment details.", "Session confirmation, pairing confirmation, live cost estimates, real-time queue states, and completion status."],
        ["Physical Printer & Node", "Local printer availability lists, system printer statuses, real-time spooler execution logs.", "Raw print buffer files, formatting options, database status sync signals, and immediate exit status updates."]
    ]

    for row_idx, row_data in enumerate(l0_data):
        row = table_l0.rows[row_idx + 1]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = l0_widths[col_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, "Calibri", 10)

    # ---------------------------------------------------------
    # SECTION 6: TECHNOLOGIES USED
    # ---------------------------------------------------------
    doc.add_page_break()
    p_sec = doc.add_paragraph()
    r_sec = p_sec.add_run("Section 6: Technologies Used")
    set_font(r_sec, "Arial", 16, COLOR_PRIMARY, bold=True)
    p_sec.paragraph_format.space_after = Pt(12)

    techs = [
        ("HTML5, CSS3, & JavaScript (ES6):", " Front-end web dashboard. Designed without heavy, complex libraries to achieve page loads under 1 second. Utilizes HTML5 HTML5 Camera API for in-browser QR scanning and standard fetch APIs for polling."),
        ("PHP 8.x (Web Server Logic):", " Primary backend server engine. Handles upload multipart form processing, files MIME inspection, session checks, and dynamic API endpoints (view_file.php, process_payment.php, get_queue.php, sync_printers.php)."),
        ("SQLite 3 via PDO:", " Relational database system. Chosen over MySQL 8.x to achieve zero-configuration deployment. The database connection falls back dynamically to the local web directory, automatically building tables on boot, making it compatible with any basic cPanel host."),
        ("Python 3.x with Tkinter (Desktop GUI Client):", " Powering the background Print Node daemon (UCPS_PrintNode.exe). Includes a custom dark-themed GUI dashboard featuring real-time spooler log panels, multithreaded polling, dynamic local printer scanning, and dynamic login validation."),
        ("SumatraPDF CLI:", " Integrated into the Windows execution client to handle PDF silent document rendering from command-line. SumatraPDF operates in the background, receiving commands without showing windows, keeping the host desktop clean."),
        ("Common UNIX Printing System (CUPS):", " Native UNIX subsystem integrated for Linux/macOS clients. The background worker dispatches print requests directly to the lp printing system with formatting constraints.")
    ]

    for title, desc in techs:
        p_tech = doc.add_paragraph(style='List Bullet')
        p_tech.paragraph_format.space_after = Pt(6)
        r_tt = p_tech.add_run(title)
        set_font(r_tt, "Calibri", 11, bold=True)
        r_td = p_tech.add_run(desc)
        set_font(r_td, "Calibri", 11)

    # ---------------------------------------------------------
    # SECTION 7: KEY FEATURES
    # ---------------------------------------------------------
    p_sec = doc.add_paragraph()
    p_sec.paragraph_format.space_before = Pt(24)
    r_sec = p_sec.add_run("Section 7: Key Features")
    set_font(r_sec, "Arial", 16, COLOR_PRIMARY, bold=True)
    p_sec.paragraph_format.space_after = Pt(12)

    features = [
        ("QR-Based Printer Pairing:", " Users scan a QR code label on a physical printer to immediately bind their browser session to that printer. A manual dropdown with status labels is also available for devices without cameras."),
        ("Real-Time SSE Print Channel:", " Instead of resource-heavy polling loops, the print node establishes a persistent Server-Sent Events (SSE) connection to stream_jobs.php, enabling instantaneous sub-second print jobs dispatch."),
        ("Interactive Operator GUI Client:", " The Windows print node runs as an executable GUI with a persistent Connection Status Bar, connection transition status messaging, and interactive printer check toggles."),
        ("Manual Printer State Control:", " Operators can check/uncheck local printers inside the client GUI, instantly triggering server-side sync to make selected printers Online or Offline manually."),
        ("SQLite Concurrency via WAL Mode:", " Activates SQLite Write-Ahead Logging (WAL) journal mode, allowing concurrent read and write operations so user operations (like logins or printer scans) are never blocked by background spoolers."),
        ("Advanced Print Configurations:", " Users can select copies count, color mode (Color vs Grayscale), paper size (A4, Letter, Legal), and page ranges dynamically on the upload card before submitting print commands."),
        ("Bypassable SumatraPDF Portable Integration:", " The desktop client checks its executable execution directory first for a portable SumatraPDF.exe binary, avoiding installation dependencies and alerting users if missing."),
        ("Live Spooler Console Mirroring:", " Critical print events (like job downloads, format conversions, and silent spooling status) are mirrored to the cloud server logs in real-time."),
        ("One-Click Web Portal Auto-Login:", " Inside the GUI client, operators can click '🌐 Open Web Portal' to open their default browser and automatically log into the web dashboard without re-typing credentials."),
        ("Dynamic Price Estimator:", " The portal counts pages using file metadata and estimates costs based on color choice and copies (5 BDT/pg for Monochrome, 15 BDT/pg for Color)."),
        ("Instant Offline Exit Notice:", " When the desktop app is closed or logged out, it sends a shutdown signal to the server, instantly marking all its printers as 'Offline' in the database. A 45-second self-healing backup check is also implemented."),
        ("Data Security & Auto-Erasure:", " The server protects uploaded documents by hashing names, checking binary MIME types, restricting folder execution using .htaccess, and automatically deleting files after printing.")
    ]

    for title, desc in features:
        p_feat = doc.add_paragraph(style='List Bullet')
        p_feat.paragraph_format.space_after = Pt(6)
        r_ft = p_feat.add_run(title)
        set_font(r_ft, "Calibri", 11, bold=True)
        r_fd = p_feat.add_run(desc)
        set_font(r_fd, "Calibri", 11)

    # ---------------------------------------------------------
    # SECTION 8: IMPLEMENTATION DETAILS
    # ---------------------------------------------------------
    doc.add_page_break()
    p_sec = doc.add_paragraph()
    r_sec = p_sec.add_run("Section 8: Implementation Details")
    set_font(r_sec, "Arial", 16, COLOR_PRIMARY, bold=True)
    p_sec.paragraph_format.space_after = Pt(12)

    p_imp = doc.add_paragraph()
    r_imp = p_imp.add_run(
        "8.1 Windows Spooler Execution Commands\n"
        "To perform silent prints on Windows, the background queue worker utilizes SumatraPDF CLI. When a PDF "
        "file is successfully downloaded from the server, the Python app executes the subprocess command:\n\n"
        "SumatraPDF.exe -print-to \"[PRINTER_NAME]\" -silent \"[LOCAL_PATH]\"\n\n"
        "If a specific printer name is not specified, it uses `-print-to-default` to spool to the default system "
        "printer. The `-silent` argument forces SumatraPDF to close immediately after transferring data to the spooler, "
        "preventing memory leaks."
    )
    set_font(r_imp, "Calibri", 11)
    p_imp.paragraph_format.space_after = Pt(12)

    p_lock = doc.add_paragraph()
    r_lock = p_lock.add_run(
        "8.2 SQLite Concurrency, WAL Mode, & Double-Checked Queue Locking\n"
        "To prevent multiple print nodes from processing the same job simultaneously while maintaining maximum web server throughput, "
        "the database uses Write-Ahead Logging (WAL) mode enabled via 'PRAGMA journal_mode = WAL;'. This allows concurrent readers (e.g. students logging in or scanning printers) "
        "and writers (e.g. local print nodes updating statuses) to access the SQLite database without lock timeouts.\n\n"
        "Additionally, the persistent streaming channel stream_jobs.php uses a read-first query pattern to avoid transaction locks when the queue is empty. "
        "It queries for pending jobs transaction-free, and starts a transaction block to write-lock the record only when a job is found:\n\n"
        "[Transaction-Free Check]\n"
        "SELECT job_id, secure_filename FROM print_jobs WHERE status = 'Pending' AND printer_id = :printer_id AND payment_status != 'Unpaid' ORDER BY upload_time ASC LIMIT 1;\n\n"
        "[If Job Found - Enter Write Lock Transaction]\n"
        "BEGIN TRANSACTION;\n"
        "SELECT status FROM print_jobs WHERE job_id = :job_id;\n"
        "UPDATE print_jobs SET status = 'Printing', processed_time = CURRENT_TIMESTAMP WHERE job_id = :job_id;\n"
        "COMMIT;\n\n"
        "This double-checked transaction ensures that each print job is printed exactly once without holding exclusive locks on the database file."
    )
    set_font(r_lock, "Calibri", 11)
    p_lock.paragraph_format.space_after = Pt(12)

    p_opts = doc.add_paragraph()
    r_opts = p_opts.add_run(
        "8.3 Dynamic Spooling Options (Copies, Color, Range, & Paper Sizing)\n"
        "When processing a print job, the spooler client parses advanced configuration options and translates them into native CLI flags for the respective operating system's spooling engine:\n\n"
        "[Windows Spooler (SumatraPDF Settings)]\n"
        "SumatraPDF is invoked with the '-print-settings' flag compiling the arguments:\n"
        "  • Copies count: '{copies}x' (e.g., '3x')\n"
        "  • Color Mode: 'color' or 'monochrome'\n"
        "  • Page Range: list of pages or ranges (e.g., '1-3,5')\n"
        "Command executed: SumatraPDF.exe -print-to \"[PRINTER]\" -print-settings \"[copies]x,[color],[range],fit\" -silent \"[PATH]\"\n\n"
        "[Linux Spooler (CUPS lp Tool Settings)]\n"
        "CUPS lp commands are constructed with corresponding parameters:\n"
        "  • Copies: '-n {copies}'\n"
        "  • Page Range: '-P {range}'\n"
        "  • Color Mode: '-o ColorModel=Color' or '-o ColorModel=Gray'\n"
        "Command executed: lp -d [PRINTER] -n [copies] -P [range] -o ColorModel=[color_val] -o fit-to-page [PATH]"
    )
    set_font(r_opts, "Calibri", 11)
    p_opts.paragraph_format.space_after = Pt(12)

    p_port = doc.add_paragraph()
    r_port = p_port.add_run(
        "8.4 Zero-Installation SumatraPDF Portable Path Checking\n"
        "To ease client client deployment on student station computers without SumatraPDF pre-installed on the system path, "
        "the spooler implements a path-search hierarchy check. It resolves the execution directory of the compiled print node "
        "binary (using os.path.dirname(sys.argv[0])) and inspects if a portable 'SumatraPDF.exe' is present in that folder first.\n\n"
        "If it is missing from both the app folder and common system paths (like Program Files or AppData), the application "
        "automatically reports detailed warning prompts to the status window, instructing the operator on where to place the portable binary."
    )
    set_font(r_port, "Calibri", 11)

    # ---------------------------------------------------------
    # SECTION 9: RESULTS & EVALUATION
    # ---------------------------------------------------------
    doc.add_page_break()
    p_sec = doc.add_paragraph()
    r_sec = p_sec.add_run("Section 9: Results & Evaluation")
    set_font(r_sec, "Arial", 16, COLOR_PRIMARY, bold=True)
    p_sec.paragraph_format.space_after = Pt(12)

    p_res = doc.add_paragraph()
    r_res = p_res.add_run(
        "The system was evaluated in a simulated campus network environment. Multiple devices (Android, iOS, "
        "Windows, macOS) uploaded various file formats concurrently. The system performed reliably under all test cases:"
    )
    set_font(r_res, "Calibri", 11)
    p_res.paragraph_format.space_after = Pt(10)

    # Test Cases Table
    table_test = doc.add_table(rows=8, cols=4)
    table_test.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_test)

    test_headers = ["Test Case", "Input/Action", "Expected Result", "Status"]
    test_widths = [Inches(1.5), Inches(1.8), Inches(2.4), Inches(0.8)]

    for idx, text in enumerate(test_headers):
        cell = table_test.rows[0].cells[idx]
        cell.width = test_widths[idx]
        set_cell_background(cell, "0D1117")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, "Arial", 10, RGBColor(255, 255, 255), bold=True)

    test_data = [
        ["Authentication Gate", "Correct/incorrect login credentials", "Approve valid sessions, reject invalid entries with error alerts", "Passed"],
        ["QR-Based Printer Pairing", "Scan printer QR code using mobile browser", "Bind user session to that printer ID in database, updating views", "Passed"],
        ["Standard PDF Print", "Upload 1.5MB PDF file via client dashboard", "Validate, hash filename, add job metadata to queue database", "Passed"],
        ["Format Restriction", "Upload malicious Windows executable (.exe)", "Block upload immediately, reject MIME check, show alert", "Passed"],
        ["FIFO Scheduling", "Upload 3 print jobs at once", "Spooler client executes jobs in the order of upload timestamps", "Passed"],
        ["Printer Offline Event", "Close Desktop Client", "Website immediately marks printers as OFFLINE and disables them", "Passed"],
        ["Immediate Cleanup", "Print job finishes", "Temp file is permanently deleted from server uploads/ folder", "Passed"]
    ]

    for row_idx, row_data in enumerate(test_data):
        row = table_test.rows[row_idx + 1]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = test_widths[col_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx < 3 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_font(r, "Calibri", 10, bold=(col_idx == 3))

    # ---------------------------------------------------------
    # QUANTITATIVE COMPARISON
    # ---------------------------------------------------------
    p_comp = doc.add_paragraph()
    p_comp.paragraph_format.space_before = Pt(24)
    r_comp = p_comp.add_run("9.2 Quantitative Workflow Efficiency Comparison")
    set_font(r_comp, "Arial", 13, COLOR_PRIMARY, bold=True)
    p_comp.paragraph_format.space_after = Pt(10)

    table_comp = doc.add_table(rows=6, cols=3)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_comp)

    comp_headers = ["Metric Evaluated", "Traditional Messenger/USB Workflow", "Universal Cloud Print System (UCPS)"]
    comp_widths = [Inches(2.0), Inches(2.2), Inches(2.3)]

    for idx, text in enumerate(comp_headers):
        cell = table_comp.rows[0].cells[idx]
        cell.width = comp_widths[idx]
        set_cell_background(cell, "0D1117")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, "Arial", 10, RGBColor(255, 255, 255), bold=True)

    comp_data = [
        ["User Steps Required", "5 steps (open chat, select file, send, wait in line, instruct operator)", "2 steps (scan QR code, upload document)"],
        ["Operator Steps Required", "6 steps (open app, download file, locate file, open file, adjust settings, print)", "1 step (ensure printer has paper/ink, monitor queue)"],
        ["Average Transfer Time", "120–180 seconds (varies based on network, download, and messaging delays)", "15–30 seconds (direct cloud upload and automatic spooling)"],
        ["Queue Visibility", "None (users must wait in a physical queue to see printer status)", "Available (real-time queue status visible on user dashboard)"],
        ["Document Privacy", "Low (files remain on public computers and operators' messaging accounts)", "High (documents are automatically deleted from server after printing)"]
    ]

    for row_idx, row_data in enumerate(comp_data):
        row = table_comp.rows[row_idx + 1]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.width = comp_widths[col_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, "Calibri", 10)

    # ---------------------------------------------------------
    # SECTION 10: CHALLENGES & LIMITATIONS
    # ---------------------------------------------------------
    doc.add_page_break()
    p_sec = doc.add_paragraph()
    r_sec = p_sec.add_run("Section 10: Challenges & Limitations")
    set_font(r_sec, "Arial", 16, COLOR_PRIMARY, bold=True)
    p_sec.paragraph_format.space_after = Pt(12)

    p_chal = doc.add_paragraph()
    p_chal.paragraph_format.line_spacing = 1.15
    r_chal = p_chal.add_run(
        "1. Dynamic Document Formatting: PDFs and images are printed with correct layouts. "
        "However, editable files (such as .docx or .pptx) can experience formatting changes if printed without Microsoft Office installed. "
        "Converting files to PDF on the client-side before upload avoids formatting shifts.\n\n"
        "2. Shared Hosting Network Constraints: Shared cPanel hosting environments often block external connections "
        "to MySQL ports. To resolve this, UCPS uses SQLite 3 database files hosted locally with CORS APIs, "
        "allowing standard HTTP requests to handle connection routing without firewall issues.\n\n"
        "3. Cache Retention Issues: Aggressive web browser caching can show printers as online even after "
        "the daemon shuts down. Implementing explicit exit signals on close (`action=offline`) and clearing URL query parameters "
        "via `window.history.replaceState` prevents browser cache issues."
    )
    set_font(r_chal, "Calibri", 11)

    # ---------------------------------------------------------
    # SECTION 11: FUTURE SCOPE
    # ---------------------------------------------------------
    p_sec = doc.add_paragraph()
    p_sec.paragraph_format.space_before = Pt(24)
    r_sec = p_sec.add_run("Section 11: Future Scope")
    set_font(r_sec, "Arial", 16, COLOR_PRIMARY, bold=True)
    p_sec.paragraph_format.space_after = Pt(12)

    p_fut = doc.add_paragraph()
    p_fut.paragraph_format.line_spacing = 1.15
    r_fut = p_fut.add_run(
        "1. Mobile App Development: Building native Android and iOS apps with built-in QR scanners will "
        "provide a faster mobile experience and support push notifications for queue updates.\n\n"
        "2. Institutional SSO Integration: Connecting the portal with university Single Sign-On (SSO) systems "
        "will allow students to use their university credentials, improving security.\n\n"
        "3. Multi-Printer Load Balancing: Developing a smart spooler algorithm that routes print jobs to the "
        "least busy printer in the lab will optimize print queue speeds.\n\n"
        "4. Raspberry Pi Print Nodes: Running the Python client daemon on low-cost Raspberry Pi microcomputers "
        "will turn standard USB printers into smart cloud printers without requiring a dedicated PC."
    )
    set_font(r_fut, "Calibri", 11)

    # ---------------------------------------------------------
    # SECTION 12: REFERENCES
    # ---------------------------------------------------------
    doc.add_page_break()
    p_sec = doc.add_paragraph()
    r_sec = p_sec.add_run("Section 12: References")
    set_font(r_sec, "Arial", 16, COLOR_PRIMARY, bold=True)
    p_sec.paragraph_format.space_after = Pt(12)

    refs = [
        "The PHP Group. 'PHP Manual: File Uploads and Filesystem Functions.' Available: https://www.php.net/manual/",
        "Mozilla Developer Network (MDN). 'Using files from web applications.' Available: https://developer.mozilla.org/",
        "Common UNIX Printing System (CUPS). 'CUPS Command-Line Printing.' Available: https://www.cups.org/documentation.html",
        "Microsoft Learn. 'Windows Printing Architecture and Spooler API.' Available: https://learn.microsoft.com/",
        "SumatraPDF. 'SumatraPDF Command Line Arguments for Silent Printing.' Available: https://www.sumatrapdfreader.org/",
        "Python Software Foundation. 'Tkinter GUI Programming for Python.' Available: https://docs.python.org/3/library/tkinter.html",
        "SQLite. 'SQLite Database File Format and Transactions.' Available: https://www.sqlite.org/",
        "SMUCT Department of Computer Science & Engineering. 'CSE Fest 2026 Software Project Showcase Rulebook.'"
    ]

    for ref in refs:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.left_indent = Inches(0.5)
        p_ref.paragraph_format.space_after = Pt(6)
        r_ref = p_ref.add_run(ref)
        set_font(r_ref, "Calibri", 11)

    # ---------------------------------------------------------
    # Header & Footer Settings
    # ---------------------------------------------------------
    # Apply headers and footers from Section 2 onwards (skipping cover page if possible,
    # but python-docx applies to all unless different_first_page_header_footer is set)
    doc.sections[0].different_first_page_header_footer = True
    
    # First page header/footer (empty)
    first_page = doc.sections[0].first_page_header
    first_page.paragraphs[0].text = ""
    first_page_f = doc.sections[0].first_page_footer
    first_page_f.paragraphs[0].text = ""
    
    # Other pages header/footer
    header = doc.sections[0].header
    p_h = header.paragraphs[0]
    p_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_h = p_h.add_run("Universal Cloud Print System (UCPS)  •  Project Report")
    set_font(r_h, "Calibri", 8.5, COLOR_MUTED)
    
    footer = doc.sections[0].footer
    p_f = footer.paragraphs[0]
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_f.add_run("Page ")
    set_font(r_f, "Calibri", 9, COLOR_MUTED)
    add_page_number(r_f)

    # Save document
    output_path = "c:\\Users\\Lenovo\\Documents\\UniversalPrinter\\UCPS_Project_Report.docx"
    doc.save(output_path)
    print(f"Document created successfully at: {output_path}")

if __name__ == "__main__":
    build_docx_report()
